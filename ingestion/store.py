# Store parsed data
# import json

# def store(data):
#     with open('data.json', 'a') as f:
#         f.write(json.dumps(data) + "\n")


import os
from docx import Document
import pandas as pd
from fpdf import FPDF

def store(data):
    os.makedirs("outputs", exist_ok=True)

    # Word
    doc = Document()
    doc.add_heading(data['title'], 0)
    for para in data['paragraphs']:
        doc.add_paragraph(para)
    doc.save("outputs/output.docx")

    # Excel
    df = pd.DataFrame({"Paragraphs": data['paragraphs'], "Links": data['links']})
    df.to_excel("outputs/output.xlsx", index=False)

    # PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, data['title'])
    for para in data['paragraphs']:
        pdf.multi_cell(0, 10, para)
    pdf.output("outputs/output.pdf")
