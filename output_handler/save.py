# output_handler/save.py
# import os
# import pandas as pd
# import textwrap
# from docx import Document
# from fpdf import FPDF
# from bs4 import BeautifulSoup

# def get_unique_output_folder(base_dir="based_output", subfolder_prefix="output"):
#     os.makedirs(base_dir, exist_ok=True)
#     count = 1
#     folder_name = os.path.join(base_dir, f"{subfolder_prefix}_{count}")
#     while os.path.exists(folder_name):
#         count += 1
#         folder_name = os.path.join(base_dir, f"{subfolder_prefix}_{count}")
#     return folder_name

# def save_output(data):
#     # print("Saving output...")   
#     # print(data)
#     folder = get_unique_output_folder()
#     os.makedirs(folder, exist_ok=True)

#     # Save raw dict to data.txt
#     with open(os.path.join(folder, "data.txt"), "w", encoding="utf-8") as f:
#         f.write(str(data))

#     # Save as TXT
#     with open(os.path.join(folder, "output.txt"), "w", encoding="utf-8") as f:
#         f.write("\n".join(data["titles"]) + "\n\n")
#         for p in data["paragraphs"]:
#             f.write(p + "\n")

#     # Save as CSV
#     pd.DataFrame(data["paragraphs"], columns=["paragraphs"]).to_csv(
#         os.path.join(folder, "output.csv"), index=False)

#     # Save as Word
#     doc = Document()
#     doc.add_heading(data["titles"], 0)
#     for p in data["paragraphs"]:
#         doc.add_paragraph(p)
#     doc.save(os.path.join(folder, "output.docx"))

#     # Save as PDF
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_left_margin(10)
#     pdf.set_right_margin(10)
#     pdf.set_font("Arial", size=12)
#     # pdf.multi_cell(0, 10, "\n".join(data["titles"]))
#     pdf.multi_cell(180, 10, "\n".join(textwrap.wrap("\n".join(data["titles"]), 100)))
    
#     # for p in data["paragraphs"]:
#     #     pdf.multi_cell(0, 10, p)
#     # pdf.output(os.path.join(folder, "output.pdf"))

#     # print(f"✅ Output saved to: {folder}")

#     for p in data["paragraphs"]:
#         wrapped = "\n".join(textwrap.wrap(p, 100))
#         pdf.multi_cell(180, 10, wrapped)

#     pdf.output(os.path.join(folder, "output.pdf"))
#     print(f"✅ Output saved to: {folder}")



# def save_output_row_text(data):
#     print("Saving output...")   
#     # print(data)
#     folder = get_unique_output_folder()
#     os.makedirs(folder, exist_ok=True)

#     # Save raw dict to data.txt
#     with open(os.path.join(folder, "data1.txt"), "w", encoding="utf-8") as f:
#         f.write(str(data))

    
#     print("ankuL")






# output_handler/save.py

# output_handler/save.py

import os
import textwrap
import pandas as pd
from docx import Document
# from fpdf import FPDF

# from fpdf import FPDF
from typing import List, Dict, Union


# --------------------------------------
# ✅ PDF Generator with Unicode Support
# --------------------------------------

# class PDFGenerator(FPDF):
#     def __init__(self):
#         super().__init__()
#         self.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
#         self.set_font('DejaVu', '', 12)
#         self.set_auto_page_break(auto=True, margin=15)


# --------------------------------------
# ✅ Output Saver Class
# --------------------------------------

class OutputSaver:
    def __init__(self, data: Dict[str, Union[str, List[str]]], base_dir: str = "based_output", subfolder_prefix: str = "output"):
        self.data = data
        self.base_dir = base_dir
        self.subfolder_prefix = subfolder_prefix
        self.output_dir = self._create_unique_folder()
        os.makedirs(self.output_dir, exist_ok=True)

    def _create_unique_folder(self) -> str:
        os.makedirs(self.base_dir, exist_ok=True)
        count = 1
        while True:
            folder = os.path.join(self.base_dir, f"{self.subfolder_prefix}_{count}")
            if not os.path.exists(folder):
                return folder
            count += 1

    def save_all_formats(self):
        self._save_raw_data()
        self._save_txt()
        self._save_csv()
        self._save_docx()
        # self._save_pdf()
        print(f"✅ Output saved to: {self.output_dir}")

    def _save_raw_data(self):
        with open(os.path.join(self.output_dir, "data.txt"), "w", encoding="utf-8") as f:
            f.write(str(self.data))

    def _save_txt(self):
        with open(os.path.join(self.output_dir, "output.txt"), "w", encoding="utf-8") as f:
            titles = self.data.get("titles", [])
            if isinstance(titles, str):
                f.write(titles + "\n\n")
            else:
                f.write("\n".join(titles) + "\n\n")

            for para in self.data.get("paragraphs", []):
                f.write(para + "\n")

    def _save_csv(self):
        paragraphs = self.data.get("paragraphs", [])
        pd.DataFrame(paragraphs, columns=["paragraphs"]).to_csv(
            os.path.join(self.output_dir, "output.csv"), index=False
        )

    def _save_docx(self):
        doc = Document()
        titles = self.data.get("titles", [])
        if isinstance(titles, str):
            doc.add_heading(titles, 0)
        elif isinstance(titles, list):
            for title in titles:
                doc.add_heading(title, 0)

        for para in self.data.get("paragraphs", []):
            doc.add_paragraph(para)

        doc.save(os.path.join(self.output_dir, "output.docx"))

    # def _save_pdf(self):
    #     pdf = PDFGenerator()
    #     pdf.add_page()

    #     titles = self.data.get("titles", [])
    #     if isinstance(titles, str):
    #         titles = [titles]
    #     for title in titles:
    #         wrapped_title = "\n".join(textwrap.wrap(str(title), width=100))
    #         pdf.multi_cell(0, 10, wrapped_title)
    #         pdf.ln()

    #     paragraphs = self.data.get("paragraphs", [])
    #     if isinstance(paragraphs, str):
    #         paragraphs = [paragraphs]
    #     for para in paragraphs:
    #         wrapped_para = "\n".join(textwrap.wrap(str(para), width=100))
    #         pdf.multi_cell(0, 10, wrapped_para)
    #         pdf.ln()

    #     try:
    #         pdf.output(os.path.join(self.output_dir, "output.pdf"))
    #     except Exception as e:
    #         print(f"[✗] Failed to save PDF: {e}")


# --------------------------------------
# ✅ Optional Raw Text Export Only
# --------------------------------------

def save_output_raw_text(data: dict):
    print("Saving raw output...")
    saver = OutputSaver(data)
    with open(os.path.join(saver.output_dir, "data1.txt"), "w", encoding="utf-8") as f:
        f.write(str(data))
    print("✅ Raw data saved.")
